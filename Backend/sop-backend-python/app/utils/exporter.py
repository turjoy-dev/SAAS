import os
import io

def generate_docx(text_content: str, title: str = "Statement of Purpose", doc_type: str = "sop", family: str = "A") -> bytes:
    """
    Generates a professionally styled Word (.docx) document tailored by document family.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header Title
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        doc.add_paragraph()

        # Family-aware formatting
        paragraphs = text_content.split("\n\n") if "\n\n" in text_content else text_content.split("\n")
        for para_text in paragraphs:
            cleaned_text = para_text.strip()
            if not cleaned_text:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(10)
            
            # Format questions or headings differently based on family
            is_heading = cleaned_text.startswith("#") or cleaned_text.startswith("Q") or cleaned_text.isupper()
            run = p.add_run(cleaned_text)
            run.font.name = "Calibri"
            
            if is_heading:
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            else:
                run.font.size = Pt(11.5)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        print("⚠️ python-docx not installed. Returning clean plain-text stream.")
        header_text = f"{title}\n{'=' * len(title)}\n\n"
        return (header_text + text_content).encode("utf-8")


def generate_pdf(text_content: str, title: str = "Statement of Purpose", doc_type: str = "sop", family: str = "A") -> bytes:
    """
    Generates a PDF document tailored by family using WeasyPrint.
    In production mode (IS_PRODUCTION=True), failures raise an exception unless CI environment is active.
    """
    try:
        from app.config import IS_PRODUCTION
    except Exception:
        IS_PRODUCTION = os.environ.get("ENVIRONMENT", "").lower() in ("production", "prod")

    html_body_paragraphs = []
    paragraphs = text_content.split("\n\n") if "\n\n" in text_content else text_content.split("\n")
    
    for p in paragraphs:
        cleaned = p.strip()
        if not cleaned:
            continue
        if cleaned.startswith("#") or cleaned.startswith("Q") or (family == "B" and "?" in cleaned):
            html_body_paragraphs.append(f"<h3 class='qa-heading'>{cleaned}</h3>")
        else:
            html_body_paragraphs.append(f"<p>{cleaned}</p>")
            
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>{title}</title>
        <style>
            @page {{ size: letter; margin: 1in; }}
            body {{ font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.6; color: #1e293b; }}
            h1 {{ text-align: center; color: #1e3a8a; font-size: 18pt; margin-bottom: 24pt; border-bottom: 2px solid #3b82f6; padding-bottom: 8pt; }}
            h3.qa-heading {{ color: #1e3a8a; font-size: 12pt; font-weight: bold; margin-top: 16pt; margin-bottom: 6pt; }}
            p {{ margin-bottom: 12pt; text-align: justify; }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        {"".join(html_body_paragraphs)}
    </body>
    </html>
    """
    try:
        from weasyprint import HTML
        return HTML(string=html_content).write_pdf()
    except Exception as e:
        if IS_PRODUCTION and not os.environ.get("CI"):
            raise RuntimeError(f"PDF generation engine failed in production: {str(e)}. PDF exports must not silently degrade to HTML.")
        print(f"⚠️ Dev/CI Mode Only: WeasyPrint fallback active ({e}). Returning styled HTML stream for testing.")
        return html_content.encode("utf-8")
